"""Generate Stage 7 I/O fixtures + goldens (deterministic, run once).

Creates in tests/fixtures/io/: text_pdf.pdf, scanned_pdf.pdf, embedded_pdf.pdf,
invoice.docx, notes.md, page.html, photo.png, and malformed files
(truncated.pdf, corrupt.docx, fake.png, unknown.bin).
Goldens (expected extracted text) live beside them as *.expected.txt.
"""
from pathlib import Path

import pymupdf
from docx import Document
from PIL import Image, ImageDraw

import sys
sys.path.insert(0, str(Path(__file__).resolve().parents[1] / "src"))
from docproc import paths  # noqa: E402

FIX = paths.TESTS_DIR / "fixtures" / "io"
FIX.mkdir(parents=True, exist_ok=True)

# ---------- PDF with a real text layer ----------
TEXT_PDF_TEXT = (
    "Invoice number 48291 issued by Acme Corporation. Total amount due is "
    "1250 dollars. Payment must be received within 30 days of the issue date."
)

doc = pymupdf.open()
page = doc.new_page(width=595, height=842)
page.insert_textbox(pymupdf.Rect(72, 72, 500, 300), TEXT_PDF_TEXT, fontsize=12)
doc.save(FIX / "text_pdf.pdf")
doc.close()
(FIX / "text_pdf.pdf.expected.txt").write_text(TEXT_PDF_TEXT, encoding="utf-8")

# ---------- scanned-style PDF: image page only, no text ----------
img = Image.new("RGB", (400, 560), "white")
d = ImageDraw.Draw(img)
for y in range(0, 560, 20):
    d.line([(0, y), (400, y)], fill=200)
img_bytes = io_png = None
import io

buf = io.BytesIO()
img.save(buf, format="PNG")

doc = pymupdf.open()
page = doc.new_page(width=595, height=842)
page.insert_image(pymupdf.Rect(100, 100, 500, 700), stream=buf.getvalue())
doc.save(FIX / "scanned_pdf.pdf")
doc.close()

# ---------- PDF with an embedded image on a text page ----------
EMBED_TEXT = (
    "Quarterly report cover page with company logo and a confidential notice "
    "stating this document is for internal distribution only."
)
logo = Image.new("RGB", (64, 64), (30, 90, 200))
lbuf = io.BytesIO()
logo.save(lbuf, format="PNG")
doc = pymupdf.open()
page = doc.new_page()
page.insert_textbox(pymupdf.Rect(50, 50, 550, 150), EMBED_TEXT, fontsize=12)
page.insert_image(pymupdf.Rect(250, 200, 350, 300), stream=lbuf.getvalue())
doc.save(FIX / "embedded_pdf.pdf")
doc.close()
(FIX / "embedded_pdf.pdf.expected.txt").write_text(EMBED_TEXT, encoding="utf-8")

# ---------- DOCX ----------
MD_TEXT = "# Meeting Minutes\n\nAttendees reviewed the budget for next quarter."
document = Document()
document.add_paragraph(MD_TEXT.splitlines()[0])
document.add_paragraph("Attendees reviewed the budget for next quarter.")
document.save(FIX / "invoice.docx")
(FIX / "invoice.docx.expected.txt").write_text(
    MD_TEXT.splitlines()[0] + "\nAttendees reviewed the budget for next quarter.",
    encoding="utf-8",
)

# ---------- Markdown ----------
md_src = "# Project Report\n\n## Findings\n\n- Accuracy improved\n- Latency reduced"
(FIX / "notes.md").write_text(md_src, encoding="utf-8")
(FIX / "notes.md.expected.txt").write_text(md_src, encoding="utf-8")

# ---------- HTML ----------
html_src = ("<html><head><title>Receipt</title></head>"
            "<body><h1>Payment Receipt</h1><p>Total: <b>99</b> USD</p></body></html>")
(FIX / "page.html").write_text(html_src, encoding="utf-8")
(FIX / "page.html.expected.txt").write_text(
    "Receipt\nPayment Receipt\nTotal:\n99\nUSD", encoding="utf-8")

# ---------- plain image ----------
photo = Image.new("RGB", (120, 160), (240, 240, 245))
ImageDraw.Draw(photo).rectangle([10, 10, 110, 150], outline=(60, 60, 60))
photo.save(FIX / "photo.png")

# ---------- malformed / unsupported ----------
(FIX / "truncated.pdf").write_bytes(b"%PDF-1.4 garbage without xref")
(FIX / "corrupt.docx").write_bytes(b"PK\x03\x04 not-a-real-zip-payload")
(FIX / "fake.png").write_bytes(b"\x89PNG\r\n\x1a\n followed by junk bytes")
(FIX / "unknown.bin").write_bytes(bytes(range(256)))

print("fixtures written to", FIX)