"""Deterministic text extraction per format; structured errors, no crashes."""
from __future__ import annotations

from pathlib import Path

from docproc.io.detect import (
    DOCX,
    HTML,
    MARKDOWN,
    PDF_TEXT,
    ParseError,
    UnsupportedFormatError,
    detect_file_type,
)


def _extract_pdf(path) -> str:
    import pymupdf

    try:
        parts = []
        with pymupdf.open(str(path)) as doc:
            for page in doc:
                parts.append(page.get_text())
        return "\n".join(parts).strip()
    except Exception as exc:
        raise ParseError("pdf", path, f"text extraction failed: {exc}") from exc


def _extract_docx(path) -> str:
    import docx

    try:
        d = docx.Document(str(path))
        lines = [p.text for p in d.paragraphs if p.text.strip()]
        for table in d.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    lines.append(" | ".join(cells))
        return "\n".join(lines)
    except Exception as exc:
        raise ParseError("docx", path, f"extraction failed: {exc}") from exc


def _extract_markdown(path) -> str:
    try:
        return Path(path).read_text(encoding="utf-8").strip()
    except UnicodeDecodeError:
        text = Path(path).read_bytes().decode("latin-1").strip()
        return text
    except Exception as exc:
        raise ParseError("md", path, f"read failed: {exc}") from exc


def _extract_html(path) -> str:
    from bs4 import BeautifulSoup

    try:
        raw = Path(path).read_bytes()
        soup = BeautifulSoup(raw, "html.parser")
        lines = [ln.strip() for ln in soup.get_text("\n").splitlines()]
        return "\n".join(ln for ln in lines if ln)
    except Exception as exc:
        raise ParseError("html", path, f"parse failed: {exc}") from exc


def extract_text(path, file_type: str | None = None) -> str:
    """Extract text by detected type; explicit file_type overrides detection."""
    p = Path(path)
    ft = file_type or detect_file_type(p).file_type
    if ft == PDF_TEXT:
        return _extract_pdf(p)
    if ft == DOCX:
        return _extract_docx(p)
    if ft == MARKDOWN:
        return _extract_markdown(p)
    if ft == HTML:
        return _extract_html(p)
    raise UnsupportedFormatError(ft, p, f"type '{ft}' has no text layer to extract")